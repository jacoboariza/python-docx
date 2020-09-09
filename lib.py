# -*- coding: utf-8 -*-
"""
Created on Wed Sep  9 23:23:52 2020

@author: Jacobo
"""

from glob import glob
from docx import Document
import docx

def listar(path, filtro=""):
    spath=path + filtro
    return glob(spath)       
    
def leer_tabla_documento (fichero):
    print ("------- Leyendo documento "+fichero)
    wordDoc = Document(fichero)

    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                print (cell.text)
    print ("------- Fin documento "+fichero)
    
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    desechar = 1
    for para in doc.paragraphs:
        if para.text=="Sección 4":
            desechar=0
        if desechar == 0:
            fullText.append(para.text)
    return '\n'.join(fullText)